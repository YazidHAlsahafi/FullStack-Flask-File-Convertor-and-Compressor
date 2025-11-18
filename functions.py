from pathlib import Path
import subprocess
from pdf2docx import Converter
from PIL import Image
from moviepy import VideoFileClip
import shutil
import platform
import fitz  # PyMuPDF
import arabic_reshaper
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import StringIO
from bidi.algorithm import get_display
import re

def find_libreoffice():
    """
    Detects LibreOffice executable depending on OS.
    Returns full path or raises FileNotFoundError.
    """
    system = platform.system()
    if system == "Windows":
        possible_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
        ]
        for path in possible_paths:
            if Path(path).exists():
                return path
        raise FileNotFoundError("LibreOffice not found. Please install or update path.")
    else:
        # Linux / Mac
        path = shutil.which("soffice") or shutil.which("libreoffice")
        if path:
            return path
        raise FileNotFoundError("LibreOffice not found in PATH. Ensure 'soffice' is installed.")

# --- Conversion Functions ---
def docx_to_pdf(input_path, output_path):
    """
    On Linux servers, the path to soffice.exe will change to just 'libreoffice'
    """
    libreoffice= find_libreoffice()
    subprocess.run([
        libreoffice, "--headless", "--convert-to", "pdf",
        "--outdir", str(Path(output_path).parent), str(input_path)
    ], check=True)


def pdf_to_docx(pdf_path, output_path):

    pdf_path = Path(pdf_path)
    output_path = Path(output_path)

    
    cv = Converter(str(pdf_path))
    # Convert all pages
    cv.convert(str(output_path), start=0, end=None)
    cv.close()


def ocr_pdf_to_docx(pdf_path, output_path, lang):
    pdf_path = Path(pdf_path)
    output_path = Path(output_path)

    # Save temp PDF in the same folder as output
    temp_pdf = output_path.parent / "ocr_temp.pdf"

    # Run OCRmyPDF command
    subprocess.run([
        "ocrmypdf",
       "--force-ocr",
        "--language", lang,
        str(pdf_path),
        str(temp_pdf)
    ], check=True)

    # Convert to DOCX
    cv = Converter(str(temp_pdf))
    cv.convert(str(output_path), start=0, end=None)
    cv.close()
    print(f"✅ OCR-based DOCX saved: {output_path}")



def pdf_to_docx_text(pdf_path, output_path):
    """
    Convert PDF to DOCX using PyMuPDF for extraction.
    Preserves page structure, applies Arabic reshaper + bidi, RTL alignment, and page numbers.
    """
    pdf_path = Path(pdf_path)
    output_path = Path(output_path)
    doc = Document()
    arabic_font = 'Times New Roman'
    font_size = 14

    pdf = fitz.open(str(pdf_path))
    for page_number in range(len(pdf)):
        page = pdf[page_number]
        page_text = page.get_text("text")  # page-level text

        # Split into lines/paragraphs
        lines = [line.strip() for line in page_text.splitlines() if line.strip()]
        for para_text in lines:
            cleaned = para_text
            try:
                reshaped = arabic_reshaper.reshape(cleaned)
            except Exception:
                reshaped = cleaned
            bidi_text = get_display(reshaped)

            paragraph = doc.add_paragraph()
            run = paragraph.add_run(bidi_text)
            run.font.name = arabic_font
            run.font.size = Pt(font_size)
            try:
                rPr = run._element.rPr
                rPr.rFonts.set(qn('w:ascii'), arabic_font)
                rPr.rFonts.set(qn('w:hAnsi'), arabic_font)
                rPr.rFonts.set(qn('w:eastAsia'), arabic_font)
            except Exception:
                pass

            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            try:
                paragraph._element.set(qn('w:bidi'), 'true')
            except Exception:
                pass

        # Add centered page number
        page_num_par = doc.add_paragraph()
        page_num_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = page_num_par.add_run(f"Page {page_number + 1}")
        run.font.name = arabic_font
        run.font.size = Pt(10)
        try:
            rPr = run._element.rPr
            rPr.rFonts.set(qn('w:ascii'), arabic_font)
            rPr.rFonts.set(qn('w:hAnsi'), arabic_font)
            rPr.rFonts.set(qn('w:eastAsia'), arabic_font)
        except Exception:
            pass

        if page_number != len(pdf) - 1:
            doc.add_page_break()

    doc.save(str(output_path))
    pdf.close()
    print(f"✅ Saved DOCX: {output_path}")


def pdf_to_text(pdf_path, output_path):
    pdf_path = Path(pdf_path)
    output_path = Path(output_path)

    pdf = fitz.open(pdf_path)
    file = open(str(output_path), 'wb')

    for page in pdf:
        text = page.get_text().encode("utf8") 
        file.write(text)
    file.close()
    



def convert_image(input_path, output_format):
    """
    Converts images between formats using Pillow.
    Works inside Docker by ensuring RGB conversion.
    """
    input_path = Path(input_path)
    output_path = input_path.with_suffix(f".{output_format.lower()}")

    try:
        img = Image.open(input_path)
        if img.mode != "RGB":
            img = img.convert("RGB")
        img.save(output_path)
        print(f"✅ Image converted to {output_path}")
        return output_path
    except Exception as e:
        raise RuntimeError(f"Image conversion failed: {e}")

def convert_video(input_path, output_format):
    """
    Converts video formats using MoviePy and FFmpeg.
    Example: MKV -> MP4, AVI -> MOV, etc.
    """
    input_path = Path(input_path)
    output_path = input_path.with_suffix(f".{output_format.lower()}")
    with VideoFileClip(str(input_path)) as clip:
        clip.write_videofile(str(output_path), codec='libx264', audio_codec='aac')
    print(f"✅ Video converted to {output_path}")
    return output_path







# --- Compression Functions ---

def compress_image(input_path, output_path, level):
    """Compress image based on user-selected level (high, medium, low)."""
    quality_map = {
        "high": 30,      # smallest file
        "medium": 60,    # balanced
        "low": 85        # best quality
    }
    quality = quality_map.get(level.lower())

    if quality is None:
        raise ValueError(f"Invalid compression level: {level}")

    img = Image.open(input_path)
    img.save(output_path, optimize=True, quality=quality)
    print(f"🖼️ Image compressed at {level} level → {output_path}")


def compress_video(input_path, output_path, level):
    """Compress video using ffmpeg based on user-selected level (high, medium, low)."""
    bitrate_map = {
        "high": "500k",     # smallest file
        "medium": "1000k",  # balanced
        "low": "2000k"      # best quality
    }
    bitrate = bitrate_map.get(level.lower())

    if bitrate is None:
        raise ValueError(f"Invalid compression level: {level}")

    subprocess.run([
        "ffmpeg", "-y",
        "-i", str(input_path),
        "-b:v", bitrate,
        str(output_path)
    ], check=True)
    print(f"🎬 Video compressed at {level} level → {output_path}")

